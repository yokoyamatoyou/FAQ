
import requests
from bs4 import BeautifulSoup
import fitz  # PyMuPDF
from docx import Document
import io

def extract_text_from_url(url):
    try:
        response = requests.get(url)
        response.raise_for_status()  # HTTPエラーをチェック
        soup = BeautifulSoup(response.text, 'html.parser')
        # スクリプトやスタイルタグを除去
        for script in soup(["script", "style"]):
            script.extract()
        text = soup.get_text()
        # 複数の空白や改行を一つにまとめる
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        return text
    except requests.exceptions.RequestException as e:
        return f"URLからのテキスト抽出エラー: {e}"

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
        return text
    except Exception as e:
        return f"PDFからのテキスト抽出エラー: {e}"

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        return f"DOCXからのテキスト抽出エラー: {e}"

# Streamlitのfile_uploaderでアップロードされたファイルオブジェクトを処理するための関数
def extract_text_from_uploaded_file(uploaded_file, file_type):
    if file_type == "pdf":
        # アップロードされたファイルのバイトデータを直接使用
        doc = fitz.open(stream=uploaded_file.getvalue(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    elif file_type == "docx":
        # アップロードされたファイルのバイトデータを直接使用
        doc = Document(io.BytesIO(uploaded_file.getvalue()))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    else:
        return "サポートされていないファイル形式です。"



