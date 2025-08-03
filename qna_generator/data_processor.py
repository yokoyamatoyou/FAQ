
import requests
from bs4 import BeautifulSoup
import fitz  # PyMuPDF
from docx import Document

MAX_UPLOAD_SIZE = 10 * 1024 * 1024  # 10MB

def extract_text_from_url(url):
    """Fetch and clean text content from the given URL.

    A 10-second timeout and a User-Agent header are used for the request.

    Raises:
        requests.exceptions.RequestException: ネットワーク関連のエラーが発生した場合。
    """
    try:
        response = requests.get(
            url,
            timeout=10,
            headers={"User-Agent": "Mozilla/5.0"},
        )
        response.raise_for_status()  # HTTPエラーをチェック
        soup = BeautifulSoup(response.text, 'html.parser')
        # スクリプトやスタイルタグを除去
        for script in soup(["script", "style"]):
            script.extract()
        text = soup.get_text()
        # 複数の空白や改行を一つにまとめる
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)
        return text
    except requests.exceptions.Timeout as e:
        raise requests.exceptions.RequestException(
            f"URLからのテキスト抽出エラー: タイムアウトが発生しました: {e}"
        ) from e
    except requests.exceptions.RequestException as e:
        raise requests.exceptions.RequestException(
            f"URLからのテキスト抽出エラー: {e}"
        ) from e

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    try:
        with fitz.open(file_path) as doc:
            text_parts = [page.get_text() for page in doc]
        return "".join(text_parts)
    except Exception as e:
        raise RuntimeError(f"PDFからのテキスト抽出エラー: {e}") from e

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        raise RuntimeError(f"DOCXからのテキスト抽出エラー: {e}") from e

# Streamlitのfile_uploaderでアップロードされたファイルオブジェクトを処理するための関数
def extract_text_from_uploaded_file(uploaded_file, file_type):
    """Handle text extraction for Streamlit-uploaded files."""
    if uploaded_file.size > MAX_UPLOAD_SIZE:
        raise ValueError("アップロードされたファイルサイズが上限(10MB)を超えています。")

    uploaded_file.seek(0)

    if file_type == "pdf":
        try:
            pdf_bytes = uploaded_file.read()
            with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
                text_parts = [page.get_text() for page in doc]
            return "".join(text_parts)
        except Exception as e:
            raise RuntimeError(f"PDFからのテキスト抽出エラー: {e}") from e
    elif file_type == "docx":
        try:
            doc = Document(uploaded_file)
            text_parts = [para.text + "\n" for para in doc.paragraphs]
            return "".join(text_parts)
        except Exception as e:
            raise RuntimeError(f"DOCXからのテキスト抽出エラー: {e}") from e
    else:
        raise ValueError("サポートされていないファイル形式です。")



